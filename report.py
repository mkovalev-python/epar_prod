import datetime
from docxtpl import DocxTemplate

def report(tree_tasks):
    from docx import Document

    def add(prefix):
        document.add_paragraph(
            el.name, style=prefix
        )
        document.add_paragraph(el.text)
        document.add_paragraph(el.text_eo)
        document.add_paragraph(el.text_ro)

    document = Document()

    for el in tree_tasks:
        if el.prefix == '1':
            add('List Number')
        if el.prefix == '11':
            add('List Number')
        if el.prefix == '111':
            add('List Number')
        if el.prefix == '1111':
            add('List Number')
        if el.prefix == '11111':
            add('List Number')

    document.add_page_break()
    src = "./tracker/media/Report_" + datetime.datetime.now().strftime("%d_%m_%Y_%H_%M") + ".docx"
    document.save(src)

    return "/protected/media/Report_" + datetime.datetime.now().strftime("%d_%m_%Y_%H_%M") + ".docx"